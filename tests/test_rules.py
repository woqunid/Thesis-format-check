from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from format_checker.rules import check_document, check_file


def test_rejects_doc_files_explicitly(tmp_path: Path):
    doc_path = tmp_path / "paper.doc"
    doc_path.write_bytes(b"not a docx")

    violations = check_file(doc_path)

    assert violations[0].problem == "暂不支持 .doc 直接检测"
    assert violations[0].expected == "请先转换为 .docx 后再检测"


def test_detects_abstract_heading_format_errors():
    document = Document()
    paragraph = document.add_paragraph("摘要")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.runs[0]
    run.font.name = "宋体"
    run.font.size = Pt(12)
    run.bold = False

    violations = check_document(document)
    problems = {item.problem for item in violations}

    assert "摘要标题字体不符合要求" in problems
    assert "摘要标题字号不符合要求" in problems
    assert "摘要标题未加粗" in problems
    assert "摘要标题未居中" in problems


def test_detects_body_paragraph_line_spacing_and_indent_errors():
    document = Document()
    heading = document.add_paragraph("1 绪论")
    heading.runs[0].font.name = "黑体"
    heading.runs[0].font.size = Pt(18)
    heading.runs[0].bold = True
    paragraph = document.add_paragraph("这是正文内容 English 123。")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(18)
    run = paragraph.runs[0]
    run.font.name = "黑体"
    run.font.size = Pt(10.5)

    violations = check_document(document)
    problems = {item.problem for item in violations}

    assert "正文中文字体不符合要求" in problems
    assert "正文行距不符合要求" in problems
    assert "正文首行缩进不符合要求" in problems
    assert "正文未两端对齐" in problems


def test_detects_keyword_separator_errors():
    document = Document()
    paragraph = document.add_paragraph("关键词：人工智能; 文档检测")
    paragraph.runs[0].font.name = "宋体"
    paragraph.runs[0].font.size = Pt(12)

    violations = check_document(document)

    assert any(item.problem == "中文关键词分隔符不符合要求" for item in violations)


def test_cover_text_is_not_checked_as_body():
    document = Document()
    paragraph = document.add_paragraph("本科毕业论文（设计）")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.runs[0].font.name = "仿宋_GB2312"
    paragraph.runs[0].font.size = Pt(36)

    violations = check_document(document)

    assert not any(item.original_text == "本科毕业论文（设计）" for item in violations)


def test_template_abstract_marker_is_checked_as_abstract_title():
    document = Document()
    paragraph = document.add_paragraph("摘□要")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.runs[0].font.name = "宋体"
    paragraph.runs[0].font.size = Pt(12)

    violations = check_document(document)
    problems = {item.problem for item in violations}

    assert "摘要标题字体不符合要求" in problems
    assert "摘要标题未居中" in problems


def test_space_separated_numbered_heading_is_checked_as_heading():
    document = Document()
    paragraph = document.add_paragraph("1 绪论")
    paragraph.runs[0].font.name = "宋体"
    paragraph.runs[0].font.size = Pt(12)
    paragraph.runs[0].bold = False

    violations = check_document(document)
    problems = {item.problem for item in violations}

    assert "标题字体不符合要求" in problems
    assert "标题字号不符合要求" in problems
    assert "标题未加粗" in problems


def test_heading_uses_chinese_run_for_chinese_font_check():
    document = Document()
    paragraph = document.add_paragraph()
    number = paragraph.add_run("1")
    number.font.name = "Times New Roman"
    number.font.size = Pt(18)
    title = paragraph.add_run(" 绪论")
    title.font.name = "黑体"
    title.font.size = Pt(18)
    title.bold = True

    violations = check_document(document)
    problems = {item.problem for item in violations}

    assert "标题字体不符合要求" not in problems


def test_toc_entries_are_not_checked_as_body_headings():
    document = Document()
    styles = document.styles
    styles.add_style("toc 1", 1)
    paragraph = document.add_paragraph("1□绪论\t1", style="toc 1")
    paragraph.runs[0].font.name = "Times New Roman"
    paragraph.runs[0].font.size = Pt(12)

    violations = check_document(document)

    assert not violations


def test_front_matter_long_paragraph_is_not_checked_as_body():
    document = Document()
    paragraph = document.add_paragraph("本人郑重声明：所呈交的论文是本人在导师的指导下独立进行研究所取得的研究成果。")
    paragraph.runs[0].font.name = "宋体"
    paragraph.runs[0].font.size = Pt(14)

    violations = check_document(document)

    assert not violations
