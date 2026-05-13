from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from format_checker.report import Violation

PT_TOLERANCE = 0.3
BODY_LINE_SPACING_PT = 23
BODY_SIZE_PT = 12
ABSTRACT_TITLE_SIZE_PT = 18
BODY_INDENT_PT = 24
TABLE_FIGURE_TITLE_SIZE_PT = 10.5


def check_file(path: Path) -> list[Violation]:
    if path.suffix.lower() == ".doc":
        return [
            Violation(
                location=str(path),
                original_text=path.name,
                problem="暂不支持 .doc 直接检测",
                current=".doc",
                expected="请先转换为 .docx 后再检测",
            )
        ]
    if path.suffix.lower() != ".docx":
        raise ValueError("只支持 .docx 文件检测")
    return check_document(Document(path))


def check_document(document: DocxDocument) -> list[Violation]:
    violations: list[Violation] = []
    violations.extend(_check_sections(document))
    for index, paragraph in enumerate(document.paragraphs, 1):
        text = paragraph.text.strip()
        if not text:
            continue
        violations.extend(_check_paragraph(index, paragraph, text))
    return violations


def _check_sections(document: DocxDocument) -> list[Violation]:
    violations: list[Violation] = []
    for index, section in enumerate(document.sections, 1):
        header_text = "\n".join(p.text.strip() for p in section.header.paragraphs if p.text.strip())
        if index > 1 and header_text and "湖北商贸学院本科毕业论文（设计）" not in header_text:
            violations.append(
                _violation(f"第 {index} 节页眉", header_text, "页眉文字不符合要求", header_text, "湖北商贸学院本科毕业论文（设计）")
            )
    return violations


def _check_paragraph(index: int, paragraph, text: str) -> list[Violation]:
    if text.startswith("摘要"):
        return _check_abstract_title(index, paragraph, text)
    if text.startswith("ABSTRACT"):
        return _check_english_abstract_title(index, paragraph, text)
    if text.startswith("关键词"):
        return _check_chinese_keywords(index, paragraph, text)
    if text.startswith("Key words"):
        return _check_english_keywords(index, paragraph, text)
    if _is_figure_or_table_title(text):
        return _check_figure_table_title(index, paragraph, text)
    if _is_heading(text):
        return _check_heading(index, paragraph, text)
    if _is_body_text(text):
        return _check_body(index, paragraph, text)
    return []


def _check_abstract_title(index: int, paragraph, text: str) -> list[Violation]:
    expected = "小二号黑体、居中、加粗"
    checks = [
        _font_violation(index, paragraph, text, "摘要标题字体不符合要求", "黑体", expected),
        _size_violation(index, paragraph, text, "摘要标题字号不符合要求", ABSTRACT_TITLE_SIZE_PT, expected),
        _bold_violation(index, paragraph, text, "摘要标题未加粗", True, expected),
        _alignment_violation(index, paragraph, text, "摘要标题未居中", WD_ALIGN_PARAGRAPH.CENTER, expected),
    ]
    return [item for item in checks if item]


def _check_english_abstract_title(index: int, paragraph, text: str) -> list[Violation]:
    expected = "小二号 Times New Roman、居中、加粗"
    checks = [
        _font_violation(index, paragraph, text, "英文摘要标题字体不符合要求", "Times New Roman", expected),
        _size_violation(index, paragraph, text, "英文摘要标题字号不符合要求", ABSTRACT_TITLE_SIZE_PT, expected),
        _bold_violation(index, paragraph, text, "英文摘要标题未加粗", True, expected),
        _alignment_violation(index, paragraph, text, "英文摘要标题未居中", WD_ALIGN_PARAGRAPH.CENTER, expected),
    ]
    return [item for item in checks if item]


def _check_chinese_keywords(index: int, paragraph, text: str) -> list[Violation]:
    violations: list[Violation] = []
    if ";" in text:
        violations.append(_violation(_location(index), text, "中文关键词分隔符不符合要求", "英文分号 ;", "中文分号 ；"))
    return violations


def _check_english_keywords(index: int, paragraph, text: str) -> list[Violation]:
    violations: list[Violation] = []
    if "；" in text:
        violations.append(_violation(_location(index), text, "英文关键词分隔符不符合要求", "中文分号 ；", "英文分号 ; 且分号后一个半角空格"))
    if ";" in text and "; " not in text:
        violations.append(_violation(_location(index), text, "英文关键词分号后缺少半角空格", "分号后无空格", "英文分号 ; 后一个半角空格"))
    return violations


def _check_figure_table_title(index: int, paragraph, text: str) -> list[Violation]:
    expected = "黑体五号、居中；数字和字母 Times New Roman 五号"
    checks = [
        _font_violation(index, paragraph, text, "图表标题中文字体不符合要求", "黑体", expected),
        _size_violation(index, paragraph, text, "图表标题字号不符合要求", TABLE_FIGURE_TITLE_SIZE_PT, expected),
        _alignment_violation(index, paragraph, text, "图表标题未居中", WD_ALIGN_PARAGRAPH.CENTER, expected),
    ]
    return [item for item in checks if item]


def _check_heading(index: int, paragraph, text: str) -> list[Violation]:
    level = text.split("□", 1)[0].count(".") + 1
    size = ABSTRACT_TITLE_SIZE_PT if level == 1 else 14 if level == 2 else BODY_SIZE_PT
    expected = _heading_expected(level)
    checks = [
        _font_violation(index, paragraph, text, "标题字体不符合要求", "黑体", expected),
        _size_violation(index, paragraph, text, "标题字号不符合要求", size, expected),
    ]
    if level <= 4:
        checks.append(_bold_violation(index, paragraph, text, "标题未加粗", True, expected))
    return [item for item in checks if item]


def _check_body(index: int, paragraph, text: str) -> list[Violation]:
    expected = "中文宋体小四，英文和数字 Times New Roman 小四；固定行距 23 磅；首行缩进 2 字符；两端对齐"
    checks = [
        _font_violation(index, paragraph, text, "正文中文字体不符合要求", "宋体", expected),
        _size_violation(index, paragraph, text, "正文字号不符合要求", BODY_SIZE_PT, expected),
        _line_spacing_violation(index, paragraph, text, expected),
        _indent_violation(index, paragraph, text, expected),
        _alignment_violation(index, paragraph, text, "正文未两端对齐", WD_ALIGN_PARAGRAPH.JUSTIFY, expected),
    ]
    return [item for item in checks if item]


def _font_violation(index: int, paragraph, text: str, problem: str, expected_font: str, expected: str) -> Violation | None:
    current = _first_run_font(paragraph)
    if current == expected_font:
        return None
    return _violation(_location(index), text, problem, current or "未设置", expected)


def _size_violation(index: int, paragraph, text: str, problem: str, expected_pt: float, expected: str) -> Violation | None:
    current = _first_run_size_pt(paragraph)
    if current is not None and abs(current - expected_pt) <= PT_TOLERANCE:
        return None
    return _violation(_location(index), text, problem, _format_pt(current), expected)


def _bold_violation(index: int, paragraph, text: str, problem: str, expected_bold: bool, expected: str) -> Violation | None:
    current = _first_run_bold(paragraph)
    if current is expected_bold:
        return None
    return _violation(_location(index), text, problem, "未加粗" if not current else "加粗", expected)


def _alignment_violation(index: int, paragraph, text: str, problem: str, expected_alignment, expected: str) -> Violation | None:
    if paragraph.alignment == expected_alignment:
        return None
    return _violation(_location(index), text, problem, _alignment_name(paragraph.alignment), expected)


def _line_spacing_violation(index: int, paragraph, text: str, expected: str) -> Violation | None:
    current = paragraph.paragraph_format.line_spacing
    current_pt = current.pt if hasattr(current, "pt") else None
    if current_pt is not None and abs(current_pt - BODY_LINE_SPACING_PT) <= PT_TOLERANCE:
        return None
    return _violation(_location(index), text, "正文行距不符合要求", _format_pt(current_pt), expected)


def _indent_violation(index: int, paragraph, text: str, expected: str) -> Violation | None:
    indent = paragraph.paragraph_format.first_line_indent
    current_pt = indent.pt if indent else None
    if current_pt is not None and current_pt >= BODY_INDENT_PT - PT_TOLERANCE:
        return None
    return _violation(_location(index), text, "正文首行缩进不符合要求", _format_pt(current_pt), expected)


def _is_heading(text: str) -> bool:
    prefix = text.split("□", 1)[0]
    return prefix.replace(".", "").isdigit() and len(prefix) <= 9


def _is_body_text(text: str) -> bool:
    return not text.startswith(("参考文献说明", "（", "(", "……", "..."))


def _is_figure_or_table_title(text: str) -> bool:
    return text.startswith(("图", "表")) and "□" in text


def _first_run_font(paragraph) -> str | None:
    run = _first_text_run(paragraph)
    return run.font.name if run else None


def _first_run_size_pt(paragraph) -> float | None:
    run = _first_text_run(paragraph)
    return run.font.size.pt if run and run.font.size else None


def _first_run_bold(paragraph) -> bool | None:
    run = _first_text_run(paragraph)
    return run.bold if run else None


def _first_text_run(paragraph):
    return next((run for run in paragraph.runs if run.text.strip()), None)


def _format_pt(value: float | None) -> str:
    return "未设置" if value is None else f"{value:g} 磅"


def _alignment_name(value) -> str:
    names = {
        WD_ALIGN_PARAGRAPH.LEFT: "左对齐",
        WD_ALIGN_PARAGRAPH.CENTER: "居中",
        WD_ALIGN_PARAGRAPH.RIGHT: "右对齐",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "两端对齐",
    }
    return names.get(value, "未设置")


def _heading_expected(level: int) -> str:
    if level == 1:
        return "一级标题：黑体小二号、加粗、顶格、独占一行"
    if level == 2:
        return "二级标题：黑体四号、加粗、顶格、独占一行"
    return "三级及以下标题：黑体小四，按层级要求加粗"


def _location(index: int) -> str:
    return f"第 {index} 段"


def _violation(location: str, text: str, problem: str, current: str, expected: str) -> Violation:
    return Violation(location, text[:120], problem, current, expected)
